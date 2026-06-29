#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word 解决方案文档生成器

按 word-solution-doc 规范生成符合中文解决方案排版要求的 .docx 文档。
- 仿宋小四正文，1.5 倍行距，首行缩进 2 字符
- 多级标题（一、/ 1、/ 1.1 / 1.1.1 / 1.1.1.1）
- 首页 + 目录页 + 正文，页码从目录页开始
- 表格自动调整、居中、首行加粗

用法：
    python generate.py --output out.docx --title "文档标题"
    python generate.py --input content.md --output out.docx
    python generate.py --demo --output out.docx
"""

import argparse
import os
import sys
from copy import deepcopy

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


# ============================================================
# 字体与字号常量
# ============================================================
FONT_HEADING = "黑体"
FONT_BODY = "仿宋"

SIZE_H1 = 22  # 二号
SIZE_H2 = 18  # 小二
SIZE_H3 = 16  # 三号
SIZE_H4 = 15  # 小三
SIZE_H5 = 14  # 四号
SIZE_BODY = 12  # 小四
SIZE_TABLE = 10.5  # 五号


# ============================================================
# XML / 排版工具
# ============================================================
def set_chinese_font(run, font_name, size_pt, bold=False):
    """设置 run 的中文字体、字号、加粗。"""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    # 东亚字符（中文）字体
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def set_line_spacing_1_5(paragraph):
    """设置 1.5 倍行距。"""
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def set_first_line_indent(paragraph, chars=2, font_size_pt=SIZE_BODY):
    """设置首行缩进（按字符数 * 字号 折算磅值）。"""
    paragraph.paragraph_format.first_line_indent = Pt(font_size_pt * chars)


# ============================================================
# 多级列表
# ============================================================
def add_multilevel_numbering(document):
    """
    在文档中注入自定义多级列表定义（abstractNum + num），并返回 numId。
    各级格式：
        0 -> 一、 (黑体 二号)
        1 -> 1、  (黑体 小二)
        2 -> 1.1  (黑体 三号)
        3 -> 1.1.1 (黑体 小三)
        4 -> 1.1.1.1 (黑体 四号)
    关键规则：一级序号不继承到下级，二级独立从 1 开始。
    """
    numbering_part = document.part.numbering_part
    if numbering_part is None:
        from docx.oxml.numbering import CT_Numbering
        from docx.parts.numbering import NumberingPart
        # 触发 python-docx 创建 numbering part
        numbering_part = document.part.numbering_part

    numbering = numbering_part.element

    # 选一个未被占用的 abstractNumId / numId
    existing_abstract_ids = [
        int(e.get(qn("w:abstractNumId")))
        for e in numbering.findall(qn("w:abstractNum"))
    ]
    existing_num_ids = [
        int(e.get(qn("w:numId")))
        for e in numbering.findall(qn("w:num"))
    ]
    abstract_id = (max(existing_abstract_ids) + 1) if existing_abstract_ids else 0
    num_id = (max(existing_num_ids) + 1) if existing_num_ids else 1

    abstract_xml = f'''
<w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
               w:abstractNumId="{abstract_id}">
  <w:multiLevelType w:val="multilevel"/>
  <w:lvl w:ilvl="0">
    <w:start w:val="1"/>
    <w:numFmt w:val="chineseCountingThousand"/>
    <w:lvlText w:val="%1、"/>
    <w:lvlJc w:val="left"/>
    <w:pPr>
      <w:ind w:left="0" w:firstLine="0"/>
    </w:pPr>
    <w:rPr>
      <w:rFonts w:ascii="黑体" w:hAnsi="黑体" w:eastAsia="黑体"/>
      <w:b/>
      <w:sz w:val="44"/>
    </w:rPr>
  </w:lvl>
  <w:lvl w:ilvl="1">
    <w:start w:val="1"/>
    <w:numFmt w:val="decimal"/>
    <w:lvlText w:val="%2、"/>
    <w:lvlJc w:val="left"/>
    <w:pPr>
      <w:ind w:left="0" w:firstLine="0"/>
    </w:pPr>
    <w:rPr>
      <w:rFonts w:ascii="黑体" w:hAnsi="黑体" w:eastAsia="黑体"/>
      <w:b/>
      <w:sz w:val="36"/>
    </w:rPr>
  </w:lvl>
  <w:lvl w:ilvl="2">
    <w:start w:val="1"/>
    <w:numFmt w:val="decimal"/>
    <w:lvlText w:val="%2.%3"/>
    <w:lvlJc w:val="left"/>
    <w:pPr>
      <w:ind w:left="0" w:firstLine="0"/>
    </w:pPr>
    <w:rPr>
      <w:rFonts w:ascii="黑体" w:hAnsi="黑体" w:eastAsia="黑体"/>
      <w:b/>
      <w:sz w:val="32"/>
    </w:rPr>
  </w:lvl>
  <w:lvl w:ilvl="3">
    <w:start w:val="1"/>
    <w:numFmt w:val="decimal"/>
    <w:lvlText w:val="%2.%3.%4"/>
    <w:lvlJc w:val="left"/>
    <w:pPr>
      <w:ind w:left="0" w:firstLine="0"/>
    </w:pPr>
    <w:rPr>
      <w:rFonts w:ascii="黑体" w:hAnsi="黑体" w:eastAsia="黑体"/>
      <w:b/>
      <w:sz w:val="30"/>
    </w:rPr>
  </w:lvl>
  <w:lvl w:ilvl="4">
    <w:start w:val="1"/>
    <w:numFmt w:val="decimal"/>
    <w:lvlText w:val="%2.%3.%4.%5"/>
    <w:lvlJc w:val="left"/>
    <w:pPr>
      <w:ind w:left="0" w:firstLine="0"/>
    </w:pPr>
    <w:rPr>
      <w:rFonts w:ascii="黑体" w:hAnsi="黑体" w:eastAsia="黑体"/>
      <w:b/>
      <w:sz w:val="28"/>
    </w:rPr>
  </w:lvl>
</w:abstractNum>
'''.strip()

    num_xml = f'''
<w:num xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
       w:numId="{num_id}">
  <w:abstractNumId w:val="{abstract_id}"/>
</w:num>
'''.strip()

    from lxml import etree
    abstract_el = etree.fromstring(abstract_xml)
    num_el = etree.fromstring(num_xml)

    # abstractNum 必须在 num 之前
    # 找到第一个 w:num 插入位置
    first_num = numbering.find(qn("w:num"))
    if first_num is not None:
        first_num.addprevious(abstract_el)
    else:
        numbering.append(abstract_el)
    numbering.append(num_el)

    return num_id


def apply_list_level(paragraph, num_id, ilvl):
    """给段落挂上多级列表的某个级别。"""
    pPr = paragraph._element.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), str(num_id))
    numPr.append(ilvl_el)
    numPr.append(numId_el)
    pPr.append(numPr)


# ============================================================
# 段落 / 标题 / 表格生成
# ============================================================
def add_body_paragraph(document, text, indent=True):
    """正文段落：仿宋 小四 1.5 倍行距 首行缩进 2 字符。"""
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_line_spacing_1_5(p)
    if indent:
        set_first_line_indent(p, chars=2, font_size_pt=SIZE_BODY)
    run = p.add_run(text)
    set_chinese_font(run, FONT_BODY, SIZE_BODY)
    return p


def add_heading_paragraph(document, text, level, num_id):
    """
    标题段落。
    level: 1~5
    num_id: 多级列表 id
    """
    size_map = {1: SIZE_H1, 2: SIZE_H2, 3: SIZE_H3, 4: SIZE_H4, 5: SIZE_H5}
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_line_spacing_1_5(p)
    pf = p.paragraph_format
    pf.first_line_indent = Pt(0)  # 标题不缩进
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    run = p.add_run(text)
    set_chinese_font(run, FONT_HEADING, size_map[level], bold=True)
    apply_list_level(p, num_id, ilvl=level - 1)
    return p


def set_table_borders(table):
    """给表格所有单元格添加 0.5pt 黑色实线框线。"""
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")       # 0.5pt = 4 eighths of a point
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def add_table(document, headers, rows, col_widths_cm=None):
    """
    表格：
    - 自动调整窗口大小
    - 文字居中
    - 表头加粗
    - 仿宋 五号
    """
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # 表格框线
    set_table_borders(table)

    # 自动调整 = WINDOW
    tbl_pr = table._element.tblPr
    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "autofit")
    tbl_pr.append(tbl_layout)

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(h)
        set_chinese_font(run, FONT_BODY, SIZE_TABLE, bold=True)

    # 数据行
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].add_run(str(val))
            set_chinese_font(run, FONT_BODY, SIZE_TABLE)

    return table


def add_cover_page(document, title, subtitle=None, meta=None):
    """
    首页（无页码），末位加分页符。
    meta: list of (label, value) 元组
    """
    # 顶部留白
    for _ in range(6):
        p = document.add_paragraph()
        set_line_spacing_1_5(p)

    # 主标题
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing_1_5(p)
    run = p.add_run(title)
    set_chinese_font(run, FONT_HEADING, 36, bold=True)

    if subtitle:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_line_spacing_1_5(p)
        run = p.add_run(subtitle)
        set_chinese_font(run, FONT_HEADING, 18, bold=False)

    # 留白
    for _ in range(8):
        p = document.add_paragraph()
        set_line_spacing_1_5(p)

    # 元信息
    if meta:
        for label, value in meta:
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_line_spacing_1_5(p)
            run = p.add_run(f"{label}：{value}")
            set_chinese_font(run, FONT_BODY, 14)

    document.add_page_break()


def add_toc_page(document):
    """目录页，使用 Word 域（TOC \\o "1-3"），首次打开需 F9 刷新。"""
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing_1_5(p)
    run = p.add_run("目  录")
    set_chinese_font(run, FONT_HEADING, SIZE_H1, bold=True)

    # TOC 域
    p = document.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = r'TOC \o "1-3" \h \z \u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "（请在 Word 中按 F9 刷新目录）"
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    run._element.append(t)
    run._element.append(fldChar3)

    document.add_page_break()


# ============================================================
# 页码
# ============================================================
def _set_section_page_numbering(section, start=None, fmt="decimal"):
    """
    在 sectPr 上设置页码格式。
    start: 起始页码（None = 续接）
    fmt: decimal / upperRoman / lowerRoman
    """
    sectPr = section._sectPr
    # 移除旧 pgNumType
    for old in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(old)
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:fmt"), fmt)
    if start is not None:
        pgNumType.set(qn("w:start"), str(start))
    sectPr.append(pgNumType)


def _remove_page_number_footer(section):
    """清除该 section 的页脚（首页不显示页码用）。"""
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()


def _add_page_number_footer(section):
    """在该 section 页脚加 PAGE 域，底部居中。"""
    footer = section.footer
    footer.is_linked_to_previous = False
    # 清空旧段落
    for p in list(footer.paragraphs):
        p._element.getparent().remove(p._element)
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    set_chinese_font(run, FONT_BODY, SIZE_BODY)

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE   \\* MERGEFORMAT"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "1"
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    run._element.append(t)
    run._element.append(fldChar3)


# ============================================================
# 页面设置
# ============================================================
def setup_page(document):
    """A4 + 标准边距。"""
    for section in document.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)


# ============================================================
# 默认样式
# ============================================================
def setup_default_style(document):
    """设置 Normal 默认样式：仿宋 小四。"""
    style = document.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(SIZE_BODY)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_BODY)
    rFonts.set(qn("w:ascii"), FONT_BODY)
    rFonts.set(qn("w:hAnsi"), FONT_BODY)


# ============================================================
# 文本清理
# ============================================================
import re

# 中文序号模式：一、二、三、…；阿拉伯数字+顿号/点号
_HEADING_NUMBER_RE = re.compile(
    r'^[一-鿿]+、'            # 一、二、三、…
    r'|^\d+(?:\.\d+)*(?:\.|\s+|、|．)'  # 1. / 1.1 / 1、/ 2．
    r'|^[IVXLCDM]+\.\s*'              # I. II. III.
)
_BOLD_MARKER_RE = re.compile(r'\*\*(.+?)\*\*')


def strip_heading_number(text):
    """去掉标题中已有的序号文字，避免与多级列表重复。"""
    return _HEADING_NUMBER_RE.sub('', text).strip()


def replace_bold(text):
    """将 **内容** 替换为 -内容-。"""
    return _BOLD_MARKER_RE.sub(r'-\1-', text)


# ============================================================
# Markdown 解析（轻量）
# ============================================================
def parse_markdown(md_text):
    """
    支持的语法：
      # / ## / ### / #### / #####  -> 标题
      普通段落
      --- 表格块 ---
      | 表头1 | 表头2 |
      | --- | --- |
      | 单元格 | 单元格 |
    返回 list[dict]，每项为：
      {"type": "h1"|"h2"|...|"body"|"table", ...}
    """
    lines = md_text.splitlines()
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        # 标题
        if line.lstrip().startswith("#"):
            stripped = line.lstrip()
            level = 0
            while level < len(stripped) and stripped[level] == "#":
                level += 1
            if level <= 5:
                text = stripped[level:].strip()
                text = replace_bold(text)       # ** → -
                text = strip_heading_number(text)           # 去掉原有序号
                blocks.append({"type": f"h{level}", "text": text})
                i += 1
                continue
        # 表格
        if line.lstrip().startswith("|") and i + 1 < len(lines) \
                and lines[i + 1].lstrip().startswith("|") \
                and "---" in lines[i + 1]:
            header = [replace_bold(c.strip()) for c in line.strip().strip("|").split("|")]
            i += 2  # 跳过分隔行
            rows = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                row = [replace_bold(c.strip()) for c in lines[i].strip().strip("|").split("|")]
                rows.append(row)
                i += 1
            blocks.append({"type": "table", "headers": header, "rows": rows})
            continue
        # 普通段落（合并连续行）
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() \
                and not lines[i].lstrip().startswith("#") \
                and not lines[i].lstrip().startswith("|"):
            para_lines.append(lines[i].strip())
            i += 1
        blocks.append({"type": "body", "text": replace_bold(" ".join(para_lines).strip())})
    return blocks


# ============================================================
# 主流程
# ============================================================
def build_document(title, blocks, cover_meta=None, subtitle=None, output_path="out.docx"):
    document = Document()
    setup_page(document)
    setup_default_style(document)

    # 多级列表
    num_id = add_multilevel_numbering(document)

    # 首页（无页码）
    add_cover_page(document, title, subtitle=subtitle,
                   meta=cover_meta or [("日期", "2026-06-28")])
    first_section = document.sections[0]
    _remove_page_number_footer(first_section)

    # 目录页（页码从 1 开始）
    add_toc_page(document)
    # 在 Word 里分节符会复制到新节，我们用页码字段重置
    _set_section_page_numbering(first_section, start=1)
    _add_page_number_footer(first_section)

    # 正文
    for block in blocks:
        t = block["type"]
        if t == "body":
            add_body_paragraph(document, block["text"])
        elif t.startswith("h") and t[1:].isdigit():
            level = int(t[1:])
            if 1 <= level <= 5:
                add_heading_paragraph(document, block["text"], level, num_id)
        elif t == "table":
            add_table(document, block["headers"], block["rows"])

    document.save(output_path)
    return output_path


def get_demo_blocks():
    """用于 --demo 的示例内容。"""
    return [
        {"type": "h1", "text": "项目概述"},
        {"type": "body", "text": "本项目旨在构建一个企业级的文档协作平台，支持多人实时编辑、版本管理、权限控制等核心能力。系统采用前后端分离架构，后端基于 Spring Cloud 微服务框架，前端使用 Vue 3 + TypeScript。"},
        {"type": "h2", "text": "项目背景"},
        {"type": "body", "text": "随着公司业务规模扩大，团队协作场景日益复杂，传统 Office 工具难以满足多人协作、版本追溯、权限管控等需求，急需建设一套统一的文档协作平台。"},
        {"type": "h3", "text": "业务需求"},
        {"type": "body", "text": "业务部门要求支持多人在同一文档上协作编辑，并能保存历史版本以便追溯。"},
        {"type": "h3", "text": "技术现状"},
        {"type": "body", "text": "目前团队使用 Git + Markdown 管理文档，无法满足非技术人员的协作需求。"},
        {"type": "h2", "text": "项目目标"},
        {"type": "body", "text": "本项目目标是在六个月内完成平台一期建设，覆盖文档协作、权限管理、版本控制三大核心模块。"},
        {"type": "h1", "text": "技术方案"},
        {"type": "h2", "text": "架构设计"},
        {"type": "body", "text": "整体架构分为四层：接入层、网关层、业务服务层、数据层。"},
        {"type": "h3", "text": "总体架构"},
        {"type": "body", "text": "接入层使用 Nginx 负责静态资源分发与负载均衡，网关层基于 Spring Cloud Gateway 实现统一鉴权与路由。"},
        {"type": "h4", "text": "核心模块"},
        {"type": "body", "text": "核心模块包括用户中心、文档服务、协作服务、权限服务。"},
        {"type": "table", "headers": ["模块", "技术栈", "说明"],
         "rows": [
             ["用户中心", "Spring Boot + MySQL", "用户与组织架构管理"],
             ["文档服务", "MongoDB + Redis", "文档存储与缓存"],
             ["协作服务", "WebSocket + CRDT", "实时协作与冲突合并"],
             ["权限服务", "Spring Security", "细粒度权限控制"],
         ]},
        {"type": "h2", "text": "技术选型"},
        {"type": "body", "text": "前端选型方面，经过对比 Vue 3、React、Svelte 三种方案，最终选择 Vue 3 作为主框架，理由是其在国内生态成熟、团队熟悉度高。"},
    ]


# ============================================================
# CLI
# ============================================================
def main():
    parser = argparse.ArgumentParser(description="Word 解决方案文档生成器")
    parser.add_argument("--output", "-o", required=True, help="输出 .docx 路径")
    parser.add_argument("--title", "-t", default="解决方案", help="文档主标题")
    parser.add_argument("--subtitle", default=None, help="文档副标题")
    parser.add_argument("--input", "-i", default=None, help="输入 Markdown 文件路径")
    parser.add_argument("--demo", action="store_true", help="使用内置示例内容")
    args = parser.parse_args()

    if args.demo:
        blocks = get_demo_blocks()
    elif args.input:
        if not os.path.exists(args.input):
            print(f"[错误] 输入文件不存在：{args.input}", file=sys.stderr)
            sys.exit(1)
        with open(args.input, "r", encoding="utf-8") as f:
            blocks = parse_markdown(f.read())
    else:
        print("[提示] 未提供 --input 或 --demo，生成空模板。", file=sys.stderr)
        blocks = []

    out = build_document(
        title=args.title,
        blocks=blocks,
        subtitle=args.subtitle,
        output_path=args.output,
    )
    print(f"[完成] 已生成：{out}")


if __name__ == "__main__":
    main()
